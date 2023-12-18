from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_document(text):
    # Split the text into parts
    parts = text.split('\n\n')

    # Create a new Word document
    doc = Document()

    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Process each part
    for part in parts:
        lines = part.split('\n')
        for i, line in enumerate(lines):
            # Add headings and normal text
            if i == 0:
                doc.add_heading(line, level=1)
            else:
                doc.add_paragraph(line)

            # Center-align the first line of each part (heading)
            if i == 0:
                paragraph = doc.paragraphs[-1]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save the document
    doc.save('Formatted_Document.docx')

# Your text
text = """
Parte 1: Introducción al Curso (30 minutos)

Bienvenida y presentación del instructor.
Objetivos del curso y expectativas de aprendizaje.
Descripción general de la importancia de los procesos ETL y el uso de Talend.

Parte 2: Validación de Conocimientos Previos (30 minutos)

Encuesta de conocimientos iniciales: Evaluación de la experiencia previa con Talend y los conceptos ETL.
Discusión de los resultados para adaptar el enfoque del curso a las necesidades específicas de los participantes.
Resumen de los conceptos clave que se abordarán durante el curso.

Módulo 1: Fundamentos de ETL y Gestión de Errores (Duración: 3 horas)

Introducción a los conceptos básicos de ETL.
Estrategias para la identificación y manejo de errores en los procesos ETL.
Caso práctico: Diseñar un flujo ETL básico utilizando Talend con gestión de errores. Los participantes crearán un proyecto simple, importarán datos y configurarán acciones en caso de errores.

Módulo 2: Optimización y Rendimiento en Procesos ETL (Duración: 4 horas)

Optimización del uso de memoria durante la ejecución de procesos ETL.
Configuración de parámetros y ajustes para mejorar el rendimiento.
Identificación y solución de problemas relacionados con la gestión de la memoria.
Caso práctico: Mejorar la eficiencia de un proceso ETL existente en Talend. Los participantes realizarán ajustes en la configuración y parámetros para optimizar el rendimiento.

Módulo 3: Big Data y Talend en la Nube (Duración: 4 horas)

Consideraciones y buenas prácticas para trabajar con entornos de Big Data.
Estrategias para trasladar desarrollos de Talend a entornos en la nube.
Caso práctico: Implementar un flujo ETL en un entorno de Big Data en la nube utilizando Talend. Los participantes realizarán la configuración y ejecución del flujo en un entorno cloud.

Módulo 4: Componentes Talend y Joblets (Duración: 4 horas)

Exploración detallada de los componentes Talend para el manejo de datos.
Uso de joblets para reutilización de lógica en diferentes jobs.
Caso práctico: Crear y utilizar joblets en un proyecto Talend. Los participantes construirán joblets y los integrarán en un flujo de trabajo más amplio.

Módulo 5: Organización y Variables en Talend (Duración: 4 horas)

Estrategias para organizar y gestionar bibliotecas de joblets.
Uso efectivo de variables de contexto para la parametrización de jobs.
Caso práctico: Configurar variables para adaptar el comportamiento del job según el entorno. Los participantes modificarán y ejecutarán jobs utilizando variables de contexto.

Módulo 6: ELT en Talend y Consideraciones de Rendimiento (Duración: 4 horas)

Diferencias entre enfoques ETL y ELT en Talend.
Implementación de procesos ELT utilizando componentes específicos.
Caso práctico: Desarrollar y optimizar un proceso ELT en Talend. Los participantes diseñarán un flujo ELT y aplicarán estrategias para mejorar el rendimiento.

Módulo 7: Diseño Eficiente y Mejores Prácticas en Talend (Duración: 4 horas)

Directrices y recomendaciones para el diseño y desarrollo eficiente en Talend.
Estrategias para la organización del código y la documentación.
Caso práctico: Revisión y mejora de un proyecto Talend existente. Los participantes colaborarán en la revisión de código y aplicarán mejores prácticas.

Módulo 8: Meta Data y Gestión de Grandes Volúmenes de Datos (Duración: 4 horas)

Manejo de la Meta Data y sus opciones a nivel de Tablas y Querys y su efectividad.
Uso de variables globales y herencia entre jobs.
Manejo de grandes volúmenes de datos (50 Millones de reg y +).
Caso práctico: Gestionar grandes volúmenes de datos utilizando Talend. Los participantes trabajarán con conjuntos de datos voluminosos y aplicarán estrategias para mejorar el rendimiento.
"""

# Create the document
create_document(text)
